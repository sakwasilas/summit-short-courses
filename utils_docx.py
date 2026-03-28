import os
import re
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

UPLOAD_IMAGE_FOLDER = "static/uploads"


def parse_docx_lesson(file_path):
    doc = Document(file_path)
    os.makedirs(UPLOAD_IMAGE_FOLDER, exist_ok=True)

    rel_image_map = {}
    image_counter = 1

    for rel in doc.part.rels.values():
        if rel.reltype == RT.IMAGE:
            ext = "png"
            partname = str(rel.target_part.partname).lower()
            if "." in partname:
                ext = partname.split(".")[-1]

            filename = f"lesson_img_{image_counter}.{ext}"
            filepath = os.path.join(UPLOAD_IMAGE_FOLDER, filename)

            with open(filepath, "wb") as f:
                f.write(rel.target_part.blob)

            rel_image_map[rel.rId] = f"/static/uploads/{filename}"
            image_counter += 1

    result = {
        "introduction": {"blocks": []},
        "outline": {"blocks": []},
        "chapters": []
    }

    practical_task = {"blocks": []}

    current_section = None
    current_title = None
    current_blocks = []

    def save_section():
        nonlocal current_section, current_title, current_blocks, practical_task

        if not current_blocks:
            return

        section_data = {"blocks": current_blocks.copy()}

        if current_section == "introduction":
            result["introduction"] = section_data
        elif current_section == "outline":
            result["outline"] = section_data
        elif current_section == "chapter":
            result["chapters"].append({
                "title": current_title,
                "content": section_data
            })
        elif current_section == "practical":
            practical_task = section_data

        current_blocks = []

    def extract_images_from_run(run):
        images = []
        xml = run._element.xml
        for rid, path in rel_image_map.items():
            if rid in xml:
                images.append(path)
        return images

    for para in doc.paragraphs:
        text = para.text.strip()
        lower = text.lower()

        is_heading = (
            lower == "introduction" or
            lower == "course outline" or
            lower == "practical task" or
            re.match(r"^chapter\s+\d+", lower)
        )

        if is_heading:
            save_section()

            if lower == "introduction":
                current_section = "introduction"
                current_title = "Introduction"
            elif lower == "course outline":
                current_section = "outline"
                current_title = "Course Outline"
            elif lower == "practical task":
                current_section = "practical"
                current_title = "Practical Task"
            else:
                current_section = "chapter"
                current_title = text

            continue

        if current_section is None:
            continue

        paragraph_images = []
        for run in para.runs:
            paragraph_images.extend(extract_images_from_run(run))

        if text:
            current_blocks.append({
                "type": "text",
                "value": text
            })

        for img in paragraph_images:
            current_blocks.append({
                "type": "image",
                "value": img
            })

    save_section()

    if (
        not result["introduction"]["blocks"] and
        not result["outline"]["blocks"] and
        not result["chapters"] and
        not practical_task["blocks"]
    ):
        fallback_blocks = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                fallback_blocks.append({
                    "type": "text",
                    "value": text
                })

        result["chapters"].append({
            "title": "Lesson Content",
            "content": {"blocks": fallback_blocks}
        })

    return result, practical_task